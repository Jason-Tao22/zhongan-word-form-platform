from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


OUTPUT_DIR = Path(__file__).resolve().parent
OUTPUT_PATH = OUTPUT_DIR / "public-demo-template.docx"


def build_public_demo_docx(output_path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Pt(42)
    section.bottom_margin = Pt(42)
    section.left_margin = Pt(42)
    section.right_margin = Pt(42)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("公开演示检验意见通知书")
    title_run.bold = True
    title_run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("文件编号：PUBLIC-2026-001")

    table = document.add_table(rows=0, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    row = table.add_row().cells
    row[0].merge(row[3])
    row[0].text = "公开演示模板"

    header = table.add_row().cells
    header[0].text = "使用单位"
    header[1].text = ""
    header[2].text = "使用登记证编号"
    header[3].text = ""

    row = table.add_row().cells
    row[0].text = "设备品种（名称）"
    row[1].text = ""
    row[2].text = "检验结论意见"
    row[3].text = ""

    row = table.add_row().cells
    row[0].text = "流动作业"
    row[1].text = "○是 ○否"
    row[2].text = "现场结论"
    row[3].text = "□合格 □整改 □不合格"

    row = table.add_row().cells
    row[0].text = "整改方式"
    row[1].text = "□现场确认 □资料确认"
    row[2].text = "联系日期"
    row[3].text = "年   月   日"

    row = table.add_row().cells
    merged = row[0].merge(row[3])
    merged.text = "问题和意见："
    merged.add_paragraph("")
    merged.add_paragraph("")
    merged.add_paragraph("")

    document.add_paragraph("人工补录区域：")
    document.add_paragraph("")
    document.add_paragraph("本通知的有效期：       年   月   日止")
    document.add_paragraph("检验人员：       日期：")
    document.add_paragraph("使用单位代表：       日期：")
    document.add_paragraph(
        "注：本模板为公开演示样本，用于说明上传、审核、发布与正式录入的完整链路。"
    )

    document.save(output_path)


if __name__ == "__main__":
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    build_public_demo_docx(OUTPUT_PATH)
    print(OUTPUT_PATH)
